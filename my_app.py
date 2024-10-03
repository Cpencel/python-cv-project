from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile pictur
document.add_picture(
    'blink.jpg', 
    width=Inches(2.0)
    )

# name phone number and email address
name = input('What is your name : ')
speak('Hello ' + name + ' how are you today?')
speak('ok ' + name + ' We just need you to answer a few questions to quickly make your cv')

speak('What is your phone number')
phone_number = input('What is your phone number : ')
speak('What is your email?')
email = input('What is your email : ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

#about me
speak('So ' + name + ' Tell me about yourself.')
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself : ')
    )

# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company you have worked in: ')
from_date = input('From Date : ')
to_date = input('To Date : ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak('Describe your experience at ' + company )
experience_details = input(
    'Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# new experiences
while True:
    speak('Do you have more experiences?')
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes': 
        p = document.add_paragraph()

        company = input('Enter company : ')
        from_date = input('From Date : ')
        to_date = input('To Date : ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        speak('Describe your experience at ' + company )
        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading('Skills')
speak('So ' + name + ' What skills do you possess?')
skill = input('What skills do you possess : ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak('Do you have more skills?')
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes': 
        speak('Which other skill?')
        skill = input('Which other skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

speak('Thank you ' + name + ' Your CV has been made.')

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using CP REACH course project"

document.save('cv.docx')